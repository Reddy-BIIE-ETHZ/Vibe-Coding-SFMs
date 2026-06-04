"""Build Supplementary Table S1 as a Word document.

Layout: portrait US Letter, one section per SFM, multi-row table.
Times New Roman 8 pt body, header shading #E8E8E8, grey borders #999999.
"""
from __future__ import annotations
import csv
from pathlib import Path
from collections import defaultdict

from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).resolve().parent
CSV_PATH = HERE / "supp_table_s1_data.csv"
OUT_PATH = HERE / "SuppTable_S1_DetailedRetrieval.docx"

SFM_ORDER = ["tSFM", "eSFM", "mhcSFM", "crisprSFM", "mir-SFM", "dtSFM"]
THRESHOLD_ORDER = {
    "tSFM":      ["identity_100", "mmseqs_080", "mmseqs_060", "mmseqs_040"],
    "eSFM":      ["identity_100", "mmseqs_080", "mmseqs_060", "mmseqs_040"],
    "mhcSFM":    ["identity_100", "mmseqs_080", "mmseqs_060", "zero_shot_rare_alleles"],
    "crisprSFM": ["identity_100", "hamming_080", "hamming_060", "hamming_040"],
    "mir-SFM":   ["identity_100", "mmseqs_080"],
    "dtSFM":     ["identity_100", "mmseqs_080", "mmseqs_060", "mmseqs_040"],
}

# Pretty threshold labels
THR_LABEL = {
    "identity_100":   "identity_100",
    "mmseqs_080":     "mmseqs_080",
    "mmseqs_060":     "mmseqs_060",
    "mmseqs_040":     "mmseqs_040",
    "hamming_080":    "hamming_080",
    "hamming_060":    "hamming_060",
    "hamming_040":    "hamming_040",
    "zero_shot_rare_alleles": "zero-shot rare alleles",
}


# Group: SFM → threshold → list of (fold_label, a2t_r1, a2t_r5, a2t_r10, t2a_r1, t2a_r5, t2a_r10, note)
def load_rows():
    by_sfm_thr_fold = defaultdict(lambda: defaultdict(dict))  # [sfm][thr][fold]["ag2ab"|"ab2ag"] = (r1,r5,r10,note)
    with CSV_PATH.open() as f:
        for row in csv.DictReader(f):
            sfm = row["sfm"]; thr = row["threshold"]; fold = row["fold"]
            d = row["direction_code"]
            by_sfm_thr_fold[sfm][thr].setdefault(fold, {})[d] = (
                row["r_at_1"], row["r_at_5"], row["r_at_10"], row["note"]
            )
    return by_sfm_thr_fold


# Map (sfm, direction_code) → ("Agent→Target" or "Target→Agent")
A2T_COL = {
    "tSFM":      "ag2ab",
    "eSFM":      "ab2ag",
    "mhcSFM":    "ag2ab",
    "crisprSFM": "ag2ab",
    "mir-SFM":   "ag2ab",
    "dtSFM":     "ag2ab",
}
DIR_NAMES = {
    "tSFM":      ("TF→DNA", "DNA→TF"),
    "eSFM":      ("E→S", "S→E"),
    "mhcSFM":    ("peptide→allele", "allele→peptide"),
    "crisprSFM": ("gRNA→OT", "OT→gRNA"),
    "mir-SFM":   ("miRNA→target", "target→miRNA"),
    "dtSFM":     ("drug→target", "target→drug"),
}


# ---------------------------------------------------------------------------
# Cell helpers
def set_cell_border(cell, color="999999", sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)


def set_cell_shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def style_cell(cell, text, *, bold=False, italic=False, size=8,
               center=True, shade=None, width_dxa=None,
               font="Times New Roman"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.line_spacing = 1.1
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, line in enumerate(str(text).split("\n")):
        if i > 0:
            p.add_run("").add_break()
        r = p.add_run(line)
        r.font.name = font
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if shade:
        set_cell_shade(cell, shade)
    if width_dxa is not None:
        cell.width = Twips(width_dxa)
    set_cell_border(cell)


# ---------------------------------------------------------------------------
def build():
    doc = Document()

    # Portrait letter, narrow margins
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width  = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(0.6)
    sec.left_margin = sec.right_margin = Inches(0.6)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(8)

    # Title
    h = doc.add_paragraph()
    r = h.add_run("Supplementary Table S1 | ")
    r.bold = True; r.font.size = Pt(11)
    r2 = h.add_run("Complete pool-512 retrieval results across six SFMs.")
    r2.font.size = Pt(11)

    # Caption
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run(
        "R@1, R@5, and R@10 (in %) for each SFM across all evaluated "
        "clustering thresholds, all cross-validation folds, and both retrieval "
        "directions. Values from 100-trial pool-512 evaluation. Fold 4 is "
        "shown but excluded from reported means where the cross-validation "
        "split was degenerate (validation set empty or epoch-0 checkpoint); "
        "asterisks (*) mark such folds. Leakage verification leakage rates "
        "for each threshold are reported in main-text Table 3. "
        "Agent → Target is the forward biological query "
        "(given an agent, retrieve its binding target); "
        "Target → Agent is the inverse. eSFM per-fold values are not "
        "archived; only CL-verified cross-fold means are shown."
    ).font.size = Pt(8)

    by_sfm = load_rows()

    # Per-SFM section
    for sfm in SFM_ORDER:
        # Section header
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(sfm)
        run.bold = True; run.font.size = Pt(10)
        run2 = h.add_run(f"  ({DIR_NAMES[sfm][0]} forward / {DIR_NAMES[sfm][1]} inverse)")
        run2.italic = True; run2.font.size = Pt(8); run2.font.color.rgb = None

        # Table: columns =
        #  Threshold | Fold | A→T R@1 | R@5 | R@10 | T→A R@1 | R@5 | R@10
        # Two header rows for the merged "Agent→Target" / "Target→Agent"

        thresholds = THRESHOLD_ORDER[sfm]
        # Determine total data rows: per threshold, 5 per-fold rows + 1 mean row
        # (eSFM mean-only; mhcSFM zero-shot single fold)
        data_rows: list[tuple] = []
        for thr in thresholds:
            folds_dict = by_sfm.get(sfm, {}).get(thr, {})

            # Identify per-fold and mean rows in canonical order
            mean_key = next((k for k in folds_dict if "mean" in str(k)), None)

            # Per-fold rows: keys "0", "0 *", "1", "1 *", etc.
            per_fold = []
            for fkey in folds_dict:
                if "mean" in str(fkey): continue
                per_fold.append(fkey)
            # Sort by integer fold number, regardless of trailing "*"
            per_fold.sort(key=lambda x: int(str(x).split()[0]))

            for fkey in per_fold:
                a2t = folds_dict[fkey].get(A2T_COL[sfm], ("—","—","—",""))
                # The other direction
                other = "ab2ag" if A2T_COL[sfm] == "ag2ab" else "ag2ab"
                t2a = folds_dict[fkey].get(other, ("—","—","—",""))
                data_rows.append((thr, fkey, a2t, t2a, "fold"))
            if mean_key is not None:
                a2t = folds_dict[mean_key].get(A2T_COL[sfm], ("—","—","—",""))
                other = "ab2ag" if A2T_COL[sfm] == "ag2ab" else "ag2ab"
                t2a = folds_dict[mean_key].get(other, ("—","—","—",""))
                data_rows.append((thr, "Mean ± SD", a2t, t2a, "mean"))

        # Build table: 2 header rows + N data rows, 8 cols
        total_rows = 2 + len(data_rows)
        # Column widths (DXA). Total page content: ~7.3 inch * 1440 ≈ 10512 DXA → cap at 10260
        col_widths = [1480, 880, 1230, 1230, 1230, 1230, 1230, 1230]
        assert sum(col_widths) == 9740, sum(col_widths)
        tbl = doc.add_table(rows=total_rows, cols=8)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.autofit = False

        # Set grid
        tblGrid = OxmlElement("w:tblGrid")
        for w_ in col_widths:
            gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w_))
            tblGrid.append(gc)
        existing = tbl._tbl.find(qn("w:tblGrid"))
        if existing is not None:
            tbl._tbl.remove(existing)
        tbl._tbl.insert(list(tbl._tbl).index(tbl._tbl.find(qn("w:tblPr"))) + 1, tblGrid)

        # Header row 1
        h1 = tbl.rows[0].cells
        h1[0].merge(tbl.rows[1].cells[0]); style_cell(h1[0], "Threshold", bold=True, shade="E8E8E8", width_dxa=col_widths[0])
        h1[1].merge(tbl.rows[1].cells[1]); style_cell(h1[1], "Fold", bold=True, shade="E8E8E8", width_dxa=col_widths[1])
        # Agent→Target merged across cols 2,3,4
        h1[2].merge(h1[3]).merge(h1[4])
        style_cell(h1[2], "Agent → Target", bold=True, shade="E8E8E8")
        # Target→Agent merged across cols 5,6,7
        h1[5].merge(h1[6]).merge(h1[7])
        style_cell(h1[5], "Target → Agent", bold=True, shade="E8E8E8")

        # Header row 2: R@1 / R@5 / R@10 × 2
        h2 = tbl.rows[1].cells
        for i, lab in enumerate(["R@1", "R@5", "R@10", "R@1", "R@5", "R@10"]):
            style_cell(h2[2 + i], lab, bold=True, shade="E8E8E8",
                       width_dxa=col_widths[2 + i])

        # Data rows
        last_thr = None
        for ri, (thr, fkey, a2t, t2a, rtype) in enumerate(data_rows):
            row = tbl.rows[2 + ri].cells

            # Threshold cell — only show on the first row for that threshold
            thr_label = THR_LABEL.get(thr, thr) if thr != last_thr else ""
            style_cell(row[0], thr_label, bold=(thr != last_thr),
                       italic=False, width_dxa=col_widths[0],
                       size=8)
            last_thr = thr

            # Fold cell
            fold_text = str(fkey)
            is_mean = (rtype == "mean")
            shade = "F4F4F4" if is_mean else None
            style_cell(row[1], fold_text, bold=is_mean, shade=shade,
                       width_dxa=col_widths[1], size=8)

            # Agent→Target columns
            for ci, val in enumerate(a2t[:3]):
                style_cell(row[2 + ci], val, bold=is_mean, shade=shade,
                           width_dxa=col_widths[2 + ci], size=8)
            # Target→Agent columns
            for ci, val in enumerate(t2a[:3]):
                style_cell(row[5 + ci], val, bold=is_mean, shade=shade,
                           width_dxa=col_widths[5 + ci], size=8)

        # Table borders + width
        tblPr = tbl._tbl.find(qn("w:tblPr"))
        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top","left","bottom","right","insideH","insideV"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:color"), "999999")
            tblBorders.append(b)
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(tblBorders)

        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), str(sum(col_widths)))
        tblW.set(qn("w:type"), "dxa")
        existing_w = tblPr.find(qn("w:tblW"))
        if existing_w is not None:
            tblPr.remove(existing_w)
        tblPr.insert(0, tblW)

    # Footer note
    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(8)
    fr = foot.add_run(
        "Notes. * Fold 4 excluded from reported means due to a cross-validation "
        "split degeneracy independently identified across all six SFMs by the "
        "orthogonal AI auditor (see Supplementary Note S2). dtSFM mmseqs_080 "
        "is single-fold (folds 1–4 lack checkpoints; see main-text caption). "
        "mhcSFM zero-shot rare alleles is a 14-allele held-out evaluation. "
        "mir-SFM mmseqs_060 / mmseqs_040 are not separately reported because "
        "mature-miRNA sequence clustering yields singleton clusters at every "
        "threshold; the mmseqs_080 row IS the by-miRNA holdout. eSFM per-fold "
        "values are not archived; only the four-fold (folds 0–3) cross-fold "
        "mean ± s.d. is shown, from the CL-verified writeup."
    )
    fr.font.size = Pt(7.5); fr.italic = True

    doc.save(OUT_PATH)
    print(f"wrote: {OUT_PATH}  ({OUT_PATH.stat().st_size:,} bytes)")


if __name__ == "__main__":
    build()
